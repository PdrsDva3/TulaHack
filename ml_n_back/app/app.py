import json
import time
from datetime import datetime, timedelta
from gettext import translation

from docx import Document
from fastapi import FastAPI, HTTPException, APIRouter
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from db.db import get_point_information_by_id, get_all_points, add_point_information, create_user, \
    get_point_by_coordinates, get_user, login_user, get_statistic_container, get_statistic_container_solve, \
    get_report_in_day, add_garbage, get_garbage_information_by_id, get_statistic_solve
from db.db import get_point_information_by_id, get_all_points, add_point_information, create_user, \
    get_point_by_coordinates, get_user, login_user, get_report_today

import os
import uuid
from fastapi.responses import FileResponse

app = FastAPI()

origins = [
    "http://localhost.tiangolo.com",
    "https://localhost.tiangolo.com",
    "http://localhost",
    "http://localhost:8080",
    "http://localhost:5173",
    "http://localhost:5174",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

points_router = APIRouter()


@points_router.get("/all")
async def get_all_points_h():
    all_points = await get_all_points()
    if not all_points:
        raise HTTPException(status_code=418, detail="i am a teapot ;)")
    return all_points


@points_router.get("/{point_id}")
async def get_point_info(point_id: int):
    point = await get_point_information_by_id(point_id)
    if not point:
        raise HTTPException(status_code=404, detail="Invalid ID")
    return point


@points_router.get("/garbage/{garbage_id}")
async def get_point_info(garbage_id: int):
    garbage = await get_garbage_information_by_id(garbage_id)
    if not garbage:
        raise HTTPException(status_code=404, detail="Invalid ID")
    return garbage


from typing import Dict, Any


class ItemRequest(BaseModel):
    name: str
    data: Dict[str, Any]

    class Config:
        arbitrary_types_allowed = True


class PointData(BaseModel):
    address: str
    lat: str
    lon: str
    photo: str

    class Config:
        arbitrary_types_allowed = True


@points_router.post("/add")
async def create_point(data: PointData):
    point = data.dict()
    point1 = await get_point_by_coordinates(point["address"], point["lat"], point["lon"])
    if not point1:
        raise HTTPException(status_code=418, detail="i am a teapot ;)")

    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    from inference import predict

    photo, prediction = predict(point["photo"])
    translation = {
        "bin": "container",  # контейнер с решеткой или отверстиями
        "tank": "tank",  # очень большой контейнер
        "container": "container",  # тут и так понятно
        "place": "place",  # место, где стоят контейнеры
        "garbage": "garbage",  # мусор
        "overflow": "overflow_container",  # заполненный/переполненный контейнер
        'large': "large_garbage"  # гора мусора вне зоны контейнеров
    }
    prediction["Container"] += prediction["overflow"]
    out = {
        "place": 0,
        "container": 0,
        "overflow_container": 0,
        "tank": 0,
        "garbage": 0,
        "large_garbage": 0
    }
    for k, v in prediction.items():
        if v != 0:
            out[translation[k.lower()]] += v

    container = dict()
    for k, v in out.items():
        if v != 0:
            container[k] = v

    if out["container"] == 0 and out["tank"] == 0:
        ret = await add_garbage(point["address"],
                                point["lat"],
                                point["lon"],
                                current_time,
                                photo,
                                'have' if prediction["garbage"] or prediction["Large"] else 'solve')
    else:
        ret = await add_point_information(
            point['address'],
            point['lat'],
            point['lon'],
            point1['problems'],
            json.dumps(container),
            current_time,
            photo,
            prediction["garbage"] + prediction["Large"],
            'bad' if prediction["garbage"] or prediction["overflow"] or prediction["Large"] else ''
        )
    if not ret:
        raise HTTPException(status_code=418, detail="i am a teapot ;)")
    return "succes"


app.include_router(points_router, prefix="/point", tags=["point"])

# ======================users

user_router = APIRouter()


@user_router.get("/{user_id}")
async def get_user_h(user_id: int):
    user = await get_user(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return user


class RegUsr(BaseModel):
    email: str
    password: str
    name: str

    class Config:
        arbitrary_types_allowed = True


@user_router.post("/registration")
async def create_user_h(data: RegUsr):
    user = data.dict()
    user_id = await create_user(user["email"], user["password"], user["name"])
    if user_id is None:
        raise HTTPException(status_code=500, detail="Internal Server Error")
    return user_id[0]


class LogUsr(BaseModel):
    email: str
    password: str

    class Config:
        arbitrary_types_allowed = True


@user_router.post("/login")
async def login_user_h(data: LogUsr):
    data = data.dict()
    user_id = await login_user(data["email"], data["password"])
    if user_id is None:
        raise HTTPException(status_code=404, detail="Invalid login or password")
    return user_id


app.include_router(user_router, prefix="/user", tags=["user"])


def create_docx(data):
    doc = Document()
    doc.add_heading('Data Report Today', 0)

    for item in data:
        doc.add_paragraph(f'Name: {item.name}, Value: {item.value}')

    unique_filename = f"{uuid.uuid4().hex}_report.docx"
    file_path = os.path.join("/tmp", unique_filename)
    doc.save(file_path)

    return file_path


report_router = APIRouter()


class ReportToday(BaseModel):
    lat: str
    lon: str

    class Config:
        arbitrary_types_allowed = True


import base64
from io import BytesIO


def decode_base64_to_image(base64_string):
    return base64.b64decode(base64_string)


def add_image_to_docx(doc, base64_image_string):
    image_bytes = decode_base64_to_image(base64_image_string)

    # Используем BytesIO, чтобы работать с байтами как с файлом
    image_stream = BytesIO(image_bytes)
    from docx.shared import Inches
    # Добавляем картинку в документ
    doc.add_picture(image_stream, width=Inches(2.3), height=Inches(2.3))


c_translation = {"place": "Количество контейнерных площадок: ",
                 "container": "Количество контейнеров: ",
                 "overflow_container": "Из них переполненны: ",
                 "tank": "Количество баков: ",
                 "garbage": "Количество ненадлежащего мусора: ",
                 "large_garbage": "Количество большого ненадлежащего мусора: "
                 }

s_translation = {"in_process": "В процессе обработки",
                 "bad": "Есть проблема",
                 "see": "Всё в порядке",
                 "no_see": "Посещения КП с прошлого раза пока не было",
                 "old": "КП не посещалась более трех дней"
                 }


async def create_file(lat, lon):
    data = await get_report_today(lat, lon)
    doc = Document()
    # Добавляем заголовок
    doc.add_heading(f'Отчет по контейнерной площадке ({lat},  {lon})', 0)
    doc.add_paragraph(f"Фото до ")
    add_image_to_docx(doc, data["photo_1"])
    if data["photo_2"]:
        doc.add_paragraph(f"Фото после ")
        add_image_to_docx(doc, data["photo_2"])
    doc.add_paragraph(f"Адрес контейнерной площадки: {data['address']}")
    doc.add_paragraph("Информация по площадке:")
    for k, v in data["containers"].items():
        doc.add_paragraph(f"    {c_translation[k]} {v}")
    doc.add_paragraph(f"Количество мусора вне контейнера: {data['other_trash']}")
    doc.add_paragraph(f"Статус контейнерной площадки: {s_translation[data['status']]}")
    unique_filename = f"{uuid.uuid4().hex}_report.docx"
    os.getcwd()
    temp_file = os.path.join("temp")
    file_path = os.path.join(temp_file, unique_filename)
    doc.save(file_path)
    return file_path


@report_router.post("/today")
async def get_today(data: ReportToday):
    datas = data.dict()
    file_path = await create_file(datas['lat'], datas['lon'])
    if not file_path:
        raise HTTPException(status_code=404, detail="Can not generate file")

    return FileResponse(path=file_path, filename='Отчет_за_сегодня.docx', media_type='multipart/form-data')


class ReportPeriod(BaseModel):
    lat: str
    lon: str
    ts_1: str
    ts_2: str

    class Config:
        arbitrary_types_allowed = True


@report_router.post("/period")
async def get_period(data: ReportPeriod):
    data = data.dict()
    doc = Document()
    # Добавляем заголовок
    doc.add_heading(f'Отчет по контейнерной площадке ({data["lat"]},  {data["lon"]})', 0)
    ts1 = (datetime.strptime(data["ts_1"], "%Y-%m-%d %H:%M:%S")).date()
    ts2 = (datetime.strptime(data["ts_2"], "%Y-%m-%d %H:%M:%S")).date()

    while ts1 != ts2 + timedelta(days=1):
        doc.add_heading(f'Отчет за {ts1}', 1)
        db_data = await get_report_in_day(data["lat"], data["lon"], ts1)

        doc.add_paragraph(f"Фото до ")
        add_image_to_docx(doc, db_data["photo_1"])
        if db_data["photo_2"]:
            doc.add_paragraph(f"Фото после ")
            add_image_to_docx(doc, db_data["photo_2"])
        doc.add_paragraph(f"Адрес контейнерной площадки: {db_data['address']}")
        for k, v in db_data["containers"].items():
            doc.add_paragraph(f"    {c_translation[k]} {v}")
        doc.add_paragraph(f"Количество мусора вне контейнера: {db_data['other_trash']}")
        doc.add_paragraph(f"Статус контейнерной площадки: {s_translation[db_data['status']]}")

        ts1 += timedelta(days=1)
    unique_filename = f"{uuid.uuid4().hex}_report.docx"
    os.getcwd()
    temp_file = os.path.join("temp")
    file_path = os.path.join(temp_file, unique_filename)
    doc.save(file_path)
    if not file_path:
        raise HTTPException(status_code=404, detail="Can not generate file")

    return FileResponse(path=file_path, filename='Отчет.docx', media_type='multipart/form-data')


app.include_router(report_router, prefix="/report", tags=["report"])

# ======================statistic

statistic_router = APIRouter()
app.include_router(statistic_router, prefix="/statistic", tags=["statistic"])


class StatisticData(BaseModel):
    ts_1: str
    ts_2: str

    class Config:
        arbitrary_types_allowed = True


@statistic_router.post("/container")
async def get_statistic(data: StatisticData):
    data = data.dict()
    ts1 = (datetime.strptime(data["ts_1"], "%Y-%m-%d %H:%M:%S")).date()

    ts2 = (datetime.strptime(data["ts_2"], "%Y-%m-%d %H:%M:%S")).date()

    statics = {
        "see": 0,
        "bad": 0,
        "no_see": 0,
    }

    while ts1 != ts2 + timedelta(days=1):
        static = await get_statistic_container(ts1)
        statics["see"] += static["see"]
        statics["bad"] += static["bad"]
        statics["no_see"] += static["no_see"]
        ts1 += timedelta(days=1)

    return statics


@statistic_router.post("/solve")
async def get_statistic(data: StatisticData):
    data = data.dict()
    ts1 = (datetime.strptime(data["ts_1"], "%Y-%m-%d %H:%M:%S")).date()

    ts2 = (datetime.strptime(data["ts_2"], "%Y-%m-%d %H:%M:%S")).date()

    statics = {
        "error": 0,
        "solve": 0,
    }

    while ts1 != ts2 + timedelta(days=1):
        static = await get_statistic_container_solve(ts1)
        statics["error"] += static["error"]
        statics["solve"] += static["solve"]
        ts1 += timedelta(days=1)

    return statics


@statistic_router.post("/trash")
async def get_statistic():
    ts1 = datetime.now().date()
    static = await get_statistic_solve(ts1)
    return static


app.include_router(statistic_router, prefix="/statistic", tags=["statistic"])
